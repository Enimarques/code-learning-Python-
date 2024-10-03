from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Criar um novo documento Word
doc = Document()

# Definir margens
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Capa
doc.add_heading('Proteção de Dados: Entenda Seus Direitos com a LGPD', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph('Como a Lei Geral de Proteção de Dados impacta você e o município').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph(" ")

# Página 1: O que é a LGPD?
doc.add_heading('O que é a LGPD?', level=2)
doc.add_paragraph(
    'A LGPD (Lei n.º 13.709/2018) é a legislação brasileira que regulamenta o tratamento de dados pessoais, '
    'com o objetivo de proteger os direitos fundamentais de liberdade e privacidade. '
    '\nObjetivo: Garantir transparência, segurança e controle sobre o uso dos dados pessoais.'
)

# Página 2: Princípios da LGPD
doc.add_heading('Princípios da LGPD', level=2)
doc.add_paragraph(
    'Transparência: Informar claramente como os dados são utilizados.\n'
    'Segurança: Proteger os dados contra acessos não autorizados e vazamentos.\n'
    'Finalidade: Os dados devem ser coletados para fins específicos e legítimos.\n'
    'Minimização dos Dados: Coletar apenas os dados necessários para a finalidade proposta.\n'
    'Adequação: O tratamento de dados deve ser compatível com as finalidades informadas.'
)

# Página 3: Direitos dos Titulares
doc.add_heading('Direitos dos Titulares', level=2)
doc.add_paragraph(
    'Acesso aos Dados: Direito de saber quais dados pessoais são coletados e como são tratados.\n'
    'Correção de Dados: Possibilidade de solicitar a correção de dados incompletos, inexatos ou desatualizados.\n'
    'Eliminação dos Dados: Direito de solicitar a exclusão dos dados pessoais, nos termos da lei.\n'
    'Portabilidade dos Dados: Transferir os dados a outro fornecedor de serviço ou produto.\n'
    'Revogação do Consentimento: O titular pode revogar o consentimento dado para o tratamento de seus dados a qualquer momento.'
)

# Página 4: Obrigações das Organizações
doc.add_heading('Obrigações das Organizações', level=2)
doc.add_paragraph(
    'Consentimento: Coletar o consentimento dos titulares de dados de forma clara e inequívoca.\n'
    'Tratamento de Dados: Tratar os dados de forma segura e transparente, respeitando os direitos dos titulares.\n'
    'Comunicação de Vazamento: Informar os titulares e a Autoridade Nacional de Proteção de Dados (ANPD) em caso de incidentes de segurança.\n'
    'Nomeação de um Encarregado: Designar um responsável pela proteção de dados pessoais (DPO).'
)

# Página 5: Penalidades
doc.add_heading('Penalidades', level=2)
doc.add_paragraph(
    'Multas: As penalidades por descumprimento da LGPD podem incluir multas de até 2% do faturamento da empresa, '
    'limitadas a R$ 50 milhões por infração.\n'
    'Suspensão de Atividades: Pode haver a suspensão parcial ou total das atividades relacionadas ao tratamento de dados.'
)

# Página 6: Como a Câmara Municipal está se Adequando
doc.add_heading('Como a Câmara Municipal está se Adequando', level=2)
doc.add_paragraph(
    'Ações Internas: Implementação de políticas e treinamentos para o tratamento adequado de dados pessoais.\n'
    'Canal de Contato: Disponibilização de um canal de atendimento para esclarecimento de dúvidas e exercício dos direitos dos titulares.\n'
    'Compromisso com a Transparência: A Câmara Municipal está comprometida em garantir a conformidade com a LGPD, '
    'protegendo os dados pessoais dos cidadãos.'
)

# Contracapa
doc.add_paragraph("\n\n")
doc.add_heading("Contato", level=2)
doc.add_paragraph(
    "Informações de contato da Câmara Municipal para mais detalhes.\n"
    "Website: Endereço do site para acesso a mais informações sobre a LGPD."
)

# Salvar o documento
docx_path = "/mnt/data/folder_LGPD.docx"
doc.save(docx_path)

# Convertendo para PDF
import comtypes.client
wdFormatPDF = 17

# Abrir o documento
word = comtypes.client.CreateObject('Word.Application')
doc = word.Documents.Open(docx_path)

# Salvar como PDF
pdf_path = "/mnt/data/folder_LGPD.pdf"
doc.SaveAs(pdf_path, FileFormat=wdFormatPDF)
doc.Close()
word.Quit()

docx_path, pdf_path
